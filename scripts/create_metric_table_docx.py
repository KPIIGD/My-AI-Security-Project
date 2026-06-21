from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\andyw\Desktop\My-AI-Security-Project")
OUT = ROOT / "output" / "doc" / "핵심_정량_지표_표_복사용.docx"


ROWS = [
    ["최종 구현 검증", "평가 데이터 수", "5,000건", "제출 전 품질검사에 사용한 문장 수"],
    [
        "최종 구현 검증",
        "전체 정밀도 / 재현율 / 균형 점수",
        "0.8803 / 0.9729 / 0.9243",
        "개인정보를 정확히 찾고 놓치지 않았는지 함께 본 성능",
    ],
    [
        "최종 구현 검증",
        "실제 조치 대상 정밀도 / 재현율 / 균형 점수",
        "0.9975 / 0.8991 / 0.9457",
        "실제 마스킹·차단이 필요한 정보 기준 성능",
    ],
    [
        "최종 구현 검증",
        "고위험 구조형 개인정보 탐지율",
        "1.0000",
        "주민등록번호, 카드번호, 계좌번호 등 위험 정보 탐지 성능",
    ],
    [
        "최종 구현 검증",
        "전화번호·이메일 탐지율",
        "1.0000",
        "실제 서비스에서 자주 입력되는 연락처 정보 탐지 성능",
    ],
    [
        "최종 구현 검증",
        "고위험 정보 전체 조치율",
        "0.9645",
        "위험도가 높은 정보가 실제 보호 조치까지 이어진 비율",
    ],
    [
        "최종 구현 검증",
        "한국어 조사 경계 정확도",
        "0.9989",
        "“홍길동이”에서 “홍길동”만 가리고 “이”는 남기는 처리 정확도",
    ],
    [
        "최종 구현 검증",
        "원문 개인정보 로그 노출",
        "0건",
        "이름, 전화번호 같은 실제 개인정보가 로그에 남지 않음",
    ],
    [
        "최종 구현 검증",
        "개인정보 위치 계산 오류",
        "0건",
        "원문에서 가려야 할 위치를 잘못 계산한 사례 없음",
    ],
    ["최종 구현 검증", "최종 품질검사 상태", "통과", "제출 가능한 품질 기준을 만족함"],
    [
        "외부 근거 보강",
        "외부 원천 데이터 파일 수",
        "28개 압축 파일",
        "문맥 판단 근거를 넓히기 위해 검토한 데이터 파일 수",
    ],
    [
        "외부 근거 보강",
        "외부 원천 데이터 총 용량",
        "26,465,440,451바이트",
        "약 26.5GB 규모의 원천 데이터 검토",
    ],
    [
        "외부 근거 보강",
        "승인된 개인정보 판단 사례",
        "1,999건",
        "개인정보 판단 근거로 사용할 수 있는 사례",
    ],
    [
        "외부 근거 보강",
        "어려운 비개인정보 사례",
        "832건",
        "개인정보처럼 보이지만 실제로는 아닌 사례",
    ],
]


COMPARE_ROWS = [
    ["기존 기본 구성", "80.15%", "19.85%", "가장 기본적인 방어 구성"],
    ["기존 구성 + 추가 판단", "90.96%", "9.04%", "AI 기반 추가 판단을 붙인 구성"],
    ["한국어 선행 필터 포함", "94.32%", "5.68%", "한국어 특화 방어선을 앞단에 추가한 구성"],
    ["전체 구성", "97.23%", "2.77%", "한국어 선행 필터와 추가 판단을 함께 사용한 구성"],
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=9, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_column_widths(table, widths_cm):
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def style_table(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_column_widths(table, widths_cm)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05


def add_metric_table(document):
    document.add_heading("핵심 정량 지표 요약", level=1)
    desc = document.add_paragraph(
        "아래 표는 캡스톤디자인 결과보고서에 그대로 붙여넣기 좋은 Word 표입니다. "
        "각 지표는 평가자가 의미를 바로 이해할 수 있도록 한글 항목명과 해석을 함께 적었습니다."
    )
    desc.paragraph_format.space_after = Pt(8)

    headers = ["구분", "평가 항목", "결과값", "의미"]
    table = document.add_table(rows=1, cols=len(headers))
    style_table(table, [3.0, 5.5, 4.0, 10.5])

    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], "1F4E79")

    for row_idx, values in enumerate(ROWS, start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[col_idx], value, size=8.5, align=align)
            if col_idx == 0:
                set_cell_shading(cells[col_idx], "D9EAF7" if values[0] == "최종 구현 검증" else "E2F0D9")
            elif row_idx % 2 == 0:
                set_cell_shading(cells[col_idx], "F7F7F7")


def add_compare_table(document):
    document.add_paragraph()
    document.add_heading("초기 비교 실험 요약", level=1)
    desc = document.add_paragraph(
        "아래 표는 한국어 특화 방어선을 넣었을 때 기존 방식보다 탐지율이 높아졌다는 점을 보여주는 보조 표입니다."
    )
    desc.paragraph_format.space_after = Pt(8)

    headers = ["비교 구성", "탐지율", "개인정보 통과율", "의미"]
    table = document.add_table(rows=1, cols=len(headers))
    style_table(table, [4.3, 3.0, 3.4, 12.3])

    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], "385723")

    for row_idx, values in enumerate(COMPARE_ROWS, start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[col_idx], value, size=8.5, align=align)
            if row_idx % 2 == 0:
                set_cell_shading(cells[col_idx], "F7F7F7")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    styles = document.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(9)

    add_metric_table(document)
    add_compare_table(document)

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
